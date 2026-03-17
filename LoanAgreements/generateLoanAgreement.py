import os
							 
from docx import Document

def erstat_tekst_i_dokument(doc, replacements):
    # Funktion til at erstatte tekst i afsnit
    def erstat_i_afsnit(paragraphs):
        for p in paragraphs:
            for key, value in replacements.items():
                if key in p.text:
                    p.text = p.text.replace(key, str(value))

    # 1. Erstat i almindelige tekstafsnit
    erstat_i_afsnit(doc.paragraphs)
    
    # 2. Erstat i eventuelle tabeller i skabelonen
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                erstat_i_afsnit(cell.paragraphs)

def opret_laaneaftale(laangiver_info):
    # Filnavnet på den uploadede skabelon
    skabelon_sti = "Laaneaftale_skabelon_privat_til_SKM.docx" 
    
										  
												   
	
    laangiver_fornavn = laangiver_info.get("fornavn", "")
    laangiver_efternavn = laangiver_info.get("efternavn", "")
    dato_indgaaelse = laangiver_info.get("dato_indgaaelse", "")
    
    # Mapning af alle pladsholdere fundet i dokumentet til de reelle værdier i dict
    laane_data = {
        "[LÅNGIVER_NAVN]": f"{laangiver_fornavn} {laangiver_efternavn}".strip(),
        "[LÅNGIVER_ADRESSE]": laangiver_info.get("adresse", ""),
        "[LÅNGIVER_CPR]": laangiver_info.get("cpr_cvr", ""),
        "[LÅNGIVER_MAIL]": laangiver_info.get("email", ""),
							
        "[BELØB]": laangiver_info.get("beloeb", ""),
        "[RENTE]": laangiver_info.get("rente", ""),
        "[DATO_INDGÅELSE]": dato_indgaaelse,
        "[DATO_BETALING]": laangiver_info.get("dato_betaling", ""),
        "[FORFALDSDATO]": laangiver_info.get("forfaldsdato", "")
    }
    
    # Tjek om skabelonen eksisterer
    if not os.path.exists(skabelon_sti):
        print(f"Fejl: Kunne ikke finde skabelonen '{skabelon_sti}'.")
        return

    # Indlæs skabelonen
    doc = Document(skabelon_sti)
    
    # Udfør udskiftning af tekst
    erstat_tekst_i_dokument(doc, laane_data)
    
    # Generer det nye filnavn: <fornavn>_<dato_indgaaelse>.docx
    # Erstatter evt. skråstreger med bindestreger i filnavnet for at undgå sti-fejl
    sikker_dato = dato_indgaaelse.replace("/", "-").replace(".", "-")
    nyt_filnavn = f"{laangiver_fornavn}_{sikker_dato}.docx"
    
    # Gem det nye dokument
    doc.save(nyt_filnavn)
    print(f"Succes! Den nye låneaftale er gemt som: {nyt_filnavn}")

# Kør scriptet
if __name__ == "__main__":
    
    # Eksempel på ordbog for ejer 1 (med dags dato som indgåelsesdato)
    ejer_1 = {
        "fornavn": "Kasper",
        "efternavn": "Warnich-Laustsen",
        "adresse": "Sjællandsgade 106, 2., 8000, Aarhus C",
        "cpr_cvr": "180594-2645",
        "email": "kasperlaustsen@gmail.com",
        "beloeb": "10.000",
        "rente": "2,0",
        "dato_indgaaelse": "10-03-2026",
        "dato_betaling": "10-03-2026",
        "forfaldsdato": "01-01-2040"
    }

    # Eksempel på ordbog for ejer 2
    ejer_2 = {
        "fornavn": "Martin",
        "efternavn": "Højlund Therkildsen",
        "adresse": "Borgmester Jakob Jensens Gade 3, 3. th, 8000 Aarhus C",
        "cpr_cvr": "270296-0897",
        "email": "martin.therkildsen@gmail.com",
        "beloeb": "10.000",
        "rente": "2,0",
        "dato_indgaaelse": "10-03-2026",
        "dato_betaling": "10-03-2026",
        "forfaldsdato": "01-01-2040"
	}								
    ejer_3 = {
        "fornavn": "Steven",
        "efternavn": "William Petersen",
        "adresse": "Harevænget 26, 6000 Kolding",
        "cpr_cvr": "230592-2989",
        "email": "swp235@gmail.com",
        "beloeb": "10.000",
        "rente": "2,0",							
        "dato_indgaaelse": "10-03-2026",
        "dato_betaling": "10-03-2026",
        "forfaldsdato": "01-01-2040"												  
	}
    
    # Kald funktionen med den ønskede ordbog
    opret_laaneaftale(ejer_1)
    opret_laaneaftale(ejer_2)
    opret_laaneaftale(ejer_3)